#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多数据库综合巡检脚本
支持: MySQL / PostgreSQL / 达梦(DM) / GoldenDB / Oracle
Python 3.8+
生成5份Word报告:
  ① 健康报告   ② 风险报告   ③ 参数异常
  ④ 空间风险   ⑤ HA风险
By 巡检工具 v1.1  2025

依赖安装:
  pip install python-docx pymysql psycopg2-binary tqdm
  # Oracle:   pip install oracledb         (推荐, 免客户端 thin 模式)
  #        或 pip install cx_Oracle        (需 Oracle Instant Client)
  # 达梦:     pip install dmPython         (需达梦客户端)
  # GoldenDB: 使用 pymysql 兼容驱动
"""
import warnings
warnings.filterwarnings("ignore")

import os
import sys
import re
import time
import datetime
import getpass
import traceback
from typing import Dict, List, Tuple, Optional, Any

# ─── 可选依赖检测 ───────────────────────────────────────────────
try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    print("提示: 未安装tqdm, 使用简单进度显示  (pip install tqdm)")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("错误: 未安装python-docx  (pip install python-docx)")
    sys.exit(1)

# ─── 报告分类映射 ────────────────────────────────────────────────
SECTION_CATEGORY_MAP: Dict[str, str] = {
    # 通用
    "HEALTH":         "健康报告",
    "PERF":           "健康报告",
    "BACKUP":         "健康报告",   # Oracle RMAN备份状态
    "STAT":           "风险报告",   # Oracle 统计信息陈旧
    "RISK":           "风险报告",
    "RISK-DEEP":      "风险报告",
    "DEEP-RISK":      "风险报告",   # Oracle 深度风险
    "PARAMETER":      "参数异常",
    "PARAMETER-RISK": "参数异常",
    "SPACE":          "空间风险",
    "SPACE-DEEP":     "空间风险",
    "HA":             "HA风险",
    "HA-DEEP":        "HA风险",
    "ADG":            "HA风险",     # Oracle Data Guard
}

REPORT_TYPES = ["健康报告", "风险报告", "参数异常", "空间风险", "HA风险"]

# 报告封面颜色主题
REPORT_THEME = {
    "健康报告": ("1F7A4D", "E8F5E9"),   # 绿色
    "风险报告": ("C0392B", "FDECEA"),   # 红色
    "参数异常": ("D35400", "FEF3E2"),   # 橙色
    "空间风险": ("1565C0", "E3F2FD"),   # 蓝色
    "HA风险":   ("6A1B9A", "F3E5F5"),   # 紫色
}

REPORT_ICON = {
    "健康报告": "✅",
    "风险报告": "⚠️",
    "参数异常": "🔧",
    "空间风险": "💾",
    "HA风险":   "🔗",
}

# ─── 进度跟踪器 ──────────────────────────────────────────────────
class ProgressTracker:
    def __init__(self):
        self.start_time = time.time()
        self.steps: Dict[str, Dict] = {}
        self.current = None

    def start_step(self, name: str, total: int = None):
        self.current = name
        self.steps[name] = {"start": time.time(), "end": None}
        if TQDM_AVAILABLE and total:
            return tqdm(total=total, desc=name, unit="项", ncols=70)
        print(f"\n▶ {name} ...")
        return None

    @staticmethod
    def tick(pbar, msg: str = ""):
        if pbar and TQDM_AVAILABLE:
            pbar.update(1)
            if msg:
                pbar.set_postfix_str(msg[:40])
        elif msg:
            print(f"  · {msg}")

    def finish_step(self, pbar=None):
        if pbar and TQDM_AVAILABLE:
            pbar.close()
        if self.current:
            t = time.time() - self.steps[self.current]["start"]
            self.steps[self.current]["end"] = time.time()
            print(f"  ✓ {self.current} 完成 ({t:.2f}s)")

    def summary(self):
        total = time.time() - self.start_time
        print("\n" + "=" * 55)
        print("  巡检耗时汇总")
        print("=" * 55)
        for k, v in self.steps.items():
            if v["end"]:
                d = v["end"] - v["start"]
                pct = d / total * 100 if total else 0
                print(f"  {k[:30]:<32} {d:6.2f}s  {pct:4.1f}%")
        print(f"  {'总计':<32} {total:6.2f}s  100.0%")
        print("=" * 55)


# ─── SQL 文件解析器 ──────────────────────────────────────────────
class SqlSectionParser:
    """将巡检SQL文件解析为带分类标签的节（Section）列表"""

    # Oracle / sqlplus 格式化命令，SQL拆分时跳过
    _ORACLE_FMT_PREFIXES = (
        "SET LINESIZE", "SET PAGESIZE", "SET TIMING", "SET FEEDBACK",
        "SET VERIFY", "SET TRIMOUT", "SET TRIMSPOOL", "SET ECHO",
        "SET SERVEROUT", "COLUMN ", "SPOOL ", "WHENEVER ",
        "TTITLE", "BTITLE", "BREAK ", "COMPUTE ", "DEFINE ",
        "VARIABLE ", "REPHEADER", "REPFOOTER",
    )

    def __init__(self, sql_text: str, db_type: str):
        self.sql_text = sql_text
        self.db_type  = db_type

    def _extract_category(self, tag: str, title: str) -> str:
        """从 [TAG] 或标题文字推断报告分类"""
        if tag:
            return SECTION_CATEGORY_MAP.get(tag.upper(), "健康报告")
        title_upper = title.upper()
        for key in SECTION_CATEGORY_MAP:
            if key in title_upper:
                return SECTION_CATEGORY_MAP[key]
        return "健康报告"

    def parse(self) -> List[Dict]:
        """返回 sections 列表, 每项: {no, title, category, sqls:[...], raw}"""
        lines = self.sql_text.splitlines()
        sections = []
        current: Optional[Dict] = None
        current_lines: List[str] = []

        # 统一 section 标记正则（覆盖 MySQL/GoldenDB/DM SELECT '>>SECTION:...'
        #   PostgreSQL \echo '>>SECTION:...'
        #   DM PRINT '>>SECTION:...'
        #   Oracle PROMPT >>SECTION:...  ← 无引号）
        section_re = re.compile(
            r">>SECTION:\s*(\d+)\.\s*(?:\[([A-Z_\-]+)\]\s*)?(.*)",
            re.IGNORECASE
        )

        def save_current():
            nonlocal current, current_lines
            if current:
                raw  = "\n".join(current_lines).strip()
                sqls = self._split_sqls(raw)
                current["sqls"] = sqls
                current["raw"]  = raw
                sections.append(current)
            current_lines = []

        for line in lines:
            m = section_re.search(line)
            if m:
                save_current()
                no    = m.group(1)
                tag   = (m.group(2) or "").strip()
                title = m.group(3).strip()
                # 去除 PROMPT / SELECT / \echo / PRINT 等前缀残留字符
                title = re.sub(r"""^[\'\"/\\]+|[\'\"/\\;]+$""", "", title).strip()
                cat   = self._extract_category(tag, title)
                current = {
                    "no":       int(no),
                    "tag":      tag,
                    "title":    title,
                    "category": cat,
                }
            else:
                current_lines.append(line)

        save_current()
        return sections

    def _split_sqls(self, raw: str) -> List[str]:
        """拆分多条SQL语句，跳过注释和格式化命令"""
        lines = []
        for line in raw.splitlines():
            stripped = line.strip()
            # 跳过整行注释
            if stripped.startswith("--") or stripped.startswith("#"):
                continue
            # 跳过 Oracle sqlplus 格式化命令
            upper_s = stripped.upper()
            if any(upper_s.startswith(p) for p in self._ORACLE_FMT_PREFIXES):
                continue
            # 跳过 psql 元命令
            if stripped.startswith("\\") and not stripped.startswith("\\G"):
                continue
            # 跳过 DM PRINT / Oracle PROMPT
            if upper_s.startswith("PRINT '") or upper_s.startswith("PROMPT "):
                if upper_s == "PROMPT":
                    continue
                continue
            lines.append(line)

        text = "\n".join(lines)
        # 按分号或 Oracle 的独立 / 行分割
        text = re.sub(r"/\s*\n", ";\n", text)
        text = re.sub(r"/\s*$", ";",    text)
        parts = re.split(r";\s*\n|;\s*$", text)
        result = []
        for p in parts:
            cleaned = self._strip_comments(p.strip())
            if cleaned:
                result.append(cleaned)
        return result

    @staticmethod
    def _strip_comments(sql: str) -> str:
        """
        移除 SQL 中的所有注释，防止中文注释触发驱动编码错误：
          1. /* ... */  块注释（含跨行）
          2. -- ...     行内尾注释
        """
        # 1. 块注释
        sql = re.sub(r"/\*.*?\*/", " ", sql, flags=re.DOTALL)
        # 2. 行尾 -- 注释（保留换行符，避免破坏多行SQL结构）
        sql = re.sub(r"--[^\n]*", "", sql)
        # 清理多余空行
        lines = [l for l in sql.splitlines() if l.strip()]
        return "\n".join(lines).strip()


# ─── 数据库连接器基类 ────────────────────────────────────────────
class BaseConnector:
    db_type = "base"

    def __init__(self):
        self.conn = None
        self.host = ""
        self.port = 0
        self.user = ""
        self.password = ""
        self.database = ""

    def get_connection_info(self):
        raise NotImplementedError

    def connect(self) -> bool:
        raise NotImplementedError

    def execute_query(self, sql: str) -> Tuple[Optional[List[str]], Optional[List[tuple]], Optional[str]]:
        """返回 (columns, rows, error)"""
        raise NotImplementedError

    @staticmethod
    def _safe_encode_check(sql: str) -> bool:
        """检查 SQL 是否含非 ASCII 字符（驱动可能拒绝）"""
        try:
            sql.encode("ascii")
            return True
        except UnicodeEncodeError:
            return False

    @staticmethod
    def _ensure_utf8(sql: str) -> str:
        """
        确保 SQL 字符串可被 UTF-8 处理。
        对含非ASCII字符的SQL，做一次 encode/decode 往返，
        消除隐藏的非UTF-8字节（如文件以 latin-1 读入的脏数据）。
        """
        try:
            return sql.encode("utf-8", errors="replace").decode("utf-8")
        except Exception:
            return sql

    def close(self):
        if self.conn:
            try:
                self.conn.close()
            except Exception:
                pass

    def _input(self, prompt: str, default: str = "", secret: bool = False) -> str:
        if default:
            prompt = f"{prompt} [{default}]: "
        else:
            prompt = f"{prompt}: "
        if secret:
            val = getpass.getpass(prompt)
        else:
            val = input(prompt).strip()
        return val if val else default


# ─── MySQL 连接器 ────────────────────────────────────────────────
class MySQLConnector(BaseConnector):
    db_type = "MySQL"

    def get_connection_info(self):
        print("\n── MySQL 连接配置 ──")
        self.host = self._input("主机地址", "127.0.0.1")
        self.port = int(self._input("端口", "3306"))
        self.user = self._input("用户名", "root")
        self.password = self._input("密码", secret=True)
        self.database = self._input("数据库", "information_schema")

    def connect(self) -> bool:
        try:
            import pymysql
            self.conn = pymysql.connect(
                host=self.host, port=self.port,
                user=self.user, password=self.password,
                database=self.database,
                charset="utf8mb4",
                use_unicode=True,        # 确保收发都走 unicode
                connect_timeout=10,
                autocommit=True
            )
            print(f"  ✓ MySQL 连接成功: {self.user}@{self.host}:{self.port}")
            return True
        except ImportError:
            print("  ✗ 未安装pymysql (pip install pymysql)")
            return False
        except Exception as e:
            print(f"  ✗ MySQL 连接失败: {e}")
            return False

    def execute_query(self, sql: str):
        sql = self._ensure_utf8(sql.strip())
        if not sql or sql.upper().startswith("SET ") or sql.upper().startswith("\\"):
            return None, None, None
        # 跳过 SHOW SLAVE STATUS\G
        sql = sql.rstrip("\\G").strip()
        try:
            with self.conn.cursor() as cur:
                cur.execute(sql)
                if cur.description:
                    cols = [d[0] for d in cur.description]
                    rows = cur.fetchall()
                    return cols, rows, None
                return None, None, None
        except UnicodeEncodeError as e:
            return None, None, f"[编码错误] SQL含非ASCII字符无法发送: {e}"
        except Exception as e:
            return None, None, str(e)


# ─── PostgreSQL 连接器 ───────────────────────────────────────────
class PostgreSQLConnector(BaseConnector):
    db_type = "PostgreSQL"

    def get_connection_info(self):
        print("\n── PostgreSQL 连接配置 ──")
        self.host = self._input("主机地址", "127.0.0.1")
        self.port = int(self._input("端口", "5432"))
        self.user = self._input("用户名", "postgres")
        self.password = self._input("密码", secret=True)
        self.database = self._input("数据库", "postgres")

    def connect(self) -> bool:
        try:
            import psycopg2
            self.conn = psycopg2.connect(
                host=self.host, port=self.port,
                user=self.user, password=self.password,
                dbname=self.database,
                connect_timeout=10,
                options="-c statement_timeout=30000 -c client_encoding=UTF8"
            )
            self.conn.autocommit = True
            print(f"  ✓ PostgreSQL 连接成功: {self.user}@{self.host}:{self.port}/{self.database}")
            return True
        except ImportError:
            print("  ✗ 未安装psycopg2 (pip install psycopg2-binary)")
            return False
        except Exception as e:
            print(f"  ✗ PostgreSQL 连接失败: {e}")
            return False

    def execute_query(self, sql: str):
        sql = self._ensure_utf8(sql.strip())
        # 跳过psql元命令
        if not sql or sql.startswith("\\") or sql.startswith("\\set") or sql.startswith("\\pset"):
            return None, None, None
        try:
            with self.conn.cursor() as cur:
                cur.execute(sql)
                if cur.description:
                    cols = [d[0] for d in cur.description]
                    rows = cur.fetchall()
                    return cols, rows, None
                return None, None, None
        except UnicodeEncodeError as e:
            return None, None, f"[编码错误] SQL含非ASCII字符无法发送: {e}"
        except Exception as e:
            try:
                self.conn.rollback()
            except Exception:
                pass
            return None, None, str(e)


# ─── 达梦(DM) 连接器 ─────────────────────────────────────────────
class DaMengConnector(BaseConnector):
    db_type = "达梦(DM)"

    def get_connection_info(self):
        print("\n── 达梦(DM) 连接配置 ──")
        self.host = self._input("主机地址", "127.0.0.1")
        self.port = int(self._input("端口", "5236"))
        self.user = self._input("用户名", "SYSDBA")
        self.password = self._input("密码", secret=True)
        self.database = self._input("数据库名(Schema)", "SYSDBA")

    def connect(self) -> bool:
        # 优先尝试 dmPython，其次尝试 jaydebeapi
        try:
            import dmPython
            self.conn = dmPython.connect(
                user=self.user,
                password=self.password,
                server=self.host,
                port=self.port,
                autoCommit=True
            )
            print(f"  ✓ 达梦 连接成功(dmPython): {self.user}@{self.host}:{self.port}")
            return True
        except ImportError:
            pass
        except Exception as e:
            print(f"  ✗ 达梦 连接失败(dmPython): {e}")
            return False

        # 尝试 pyodbc
        try:
            import pyodbc
            dsn = (
                f"DRIVER={{DM8 ODBC DRIVER}};"
                f"SERVER={self.host};PORT={self.port};"
                f"UID={self.user};PWD={self.password}"
            )
            self.conn = pyodbc.connect(dsn, autocommit=True, timeout=10)
            print(f"  ✓ 达梦 连接成功(ODBC): {self.user}@{self.host}:{self.port}")
            return True
        except ImportError:
            print("  ✗ 未安装dmPython或pyodbc，请安装达梦客户端驱动")
            return False
        except Exception as e:
            print(f"  ✗ 达梦 连接失败: {e}")
            return False

    def execute_query(self, sql: str):
        sql = self._ensure_utf8(sql.strip())
        if not sql or sql.upper().startswith("SET LINESIZE") or sql.upper().startswith("SET PAGESIZE"):
            return None, None, None
        if sql.upper().startswith("PRINT "):
            return None, None, None
        # 去掉末尾分号
        sql = sql.rstrip(";").strip()
        if not sql:
            return None, None, None
        try:
            cur = self.conn.cursor()
            cur.execute(sql)
            if cur.description:
                cols = [d[0] for d in cur.description]
                rows = cur.fetchall()
                cur.close()
                return cols, rows, None
            cur.close()
            return None, None, None
        except UnicodeEncodeError as e:
            return None, None, f"[编码错误] SQL含非ASCII字符无法发送: {e}"
        except Exception as e:
            return None, None, str(e)


# ─── GoldenDB 连接器 ─────────────────────────────────────────────
class GoldenDBConnector(MySQLConnector):
    """GoldenDB 底层兼容 MySQL 协议，复用 MySQL 连接器"""
    db_type = "GoldenDB"

    def get_connection_info(self):
        print("\n── GoldenDB 连接配置 (CN协调节点) ──")
        self.host = self._input("CN节点地址", "127.0.0.1")
        self.port = int(self._input("CN端口", "3308"))
        self.user = self._input("用户名", "root")
        self.password = self._input("密码", secret=True)
        self.database = self._input("数据库", "information_schema")

    def execute_query(self, sql: str):
        sql = sql.strip()
        # GoldenDB 特有命令降级处理
        upper = sql.upper().strip()
        if any(upper.startswith(k) for k in [
            "SHOW CLUSTER", "SHOW NODES", "SHOW SHARD",
            "SET SESSION", "SET @", "\\"
        ]):
            return None, None, "GoldenDB-specific DDL (skipped in offline mode)"
        return super().execute_query(sql)


# ─── Oracle 连接器 ───────────────────────────────────────────────
class OracleConnector(BaseConnector):
    """
    Oracle 连接器
    优先使用 python-oracledb (thin 模式，免 Oracle Client 安装)
    回退使用 cx_Oracle (需要 Oracle Instant Client)
    """
    db_type = "Oracle"

    # Oracle sqlplus 格式化命令前缀，执行时跳过
    _SKIP_PREFIXES = (
        "SET ", "COLUMN ", "SPOOL ", "PROMPT", "WHENEVER ",
        "DEFINE ", "VARIABLE ", "EXEC ", "HOST ", "@", "REM ",
        "TTITLE", "BTITLE", "BREAK ", "COMPUTE ", "REPHEADER",
        "REPFOOTER", "--", "/*",
    )

    def __init__(self):
        super().__init__()
        self.service_name = ""
        self.sid          = ""
        self._driver      = None   # "oracledb" 或 "cx_Oracle"

    def get_connection_info(self):
        print("\n── Oracle 连接配置 ──")
        self.host = self._input("主机地址", "127.0.0.1")
        self.port = int(self._input("监听端口", "1521"))
        conn_type = input("  连接方式 (1-SERVICE_NAME  2-SID) [1]: ").strip() or "1"
        if conn_type == "2":
            self.sid = self._input("SID", "orcl")
        else:
            self.service_name = self._input("SERVICE_NAME", "orcl")
        self.user     = self._input("用户名", "system")
        self.password = self._input("密码", secret=True)

        # 是否以 SYSDBA 连接
        sysdba = input("  以 SYSDBA 权限连接? (y/N): ").strip().lower()
        self._as_sysdba = (sysdba == "y")

    def connect(self) -> bool:
        # 尝试 python-oracledb (thin)
        try:
            import oracledb
            self._driver = "oracledb"
            if self.service_name:
                dsn = oracledb.makedsn(self.host, self.port,
                                       service_name=self.service_name)
            else:
                dsn = oracledb.makedsn(self.host, self.port, sid=self.sid)

            kwargs = dict(user=self.user, password=self.password, dsn=dsn)
            if getattr(self, "_as_sysdba", False):
                kwargs["mode"] = oracledb.AUTH_MODE_SYSDBA
            self.conn = oracledb.connect(**kwargs)
            print(f"  ✓ Oracle 连接成功(oracledb-thin): "
                  f"{self.user}@{self.host}:{self.port}")
            return True
        except ImportError:
            pass
        except Exception as e:
            print(f"  ✗ Oracle oracledb 连接失败: {e}")
            # 继续尝试 cx_Oracle

        # 回退 cx_Oracle
        try:
            import cx_Oracle
            self._driver = "cx_Oracle"
            if self.service_name:
                dsn = cx_Oracle.makedsn(self.host, self.port,
                                        service_name=self.service_name)
            else:
                dsn = cx_Oracle.makedsn(self.host, self.port, sid=self.sid)

            kwargs = dict(user=self.user, password=self.password, dsn=dsn,
                          encoding="UTF-8", nencoding="UTF-8")   # ← 关键：避免ASCII编码错误
            if getattr(self, "_as_sysdba", False):
                kwargs["mode"] = cx_Oracle.SYSDBA
            self.conn = cx_Oracle.connect(**kwargs)
            print(f"  ✓ Oracle 连接成功(cx_Oracle): "
                  f"{self.user}@{self.host}:{self.port}")
            return True
        except ImportError:
            print("  ✗ 未安装Oracle驱动: pip install oracledb  "
                  "(或 pip install cx_Oracle + Oracle Instant Client)")
            return False
        except Exception as e:
            print(f"  ✗ Oracle cx_Oracle 连接失败: {e}")
            return False

    def execute_query(self, sql: str):
        sql = self._ensure_utf8(sql.strip())
        if not sql:
            return None, None, None

        # 跳过 sqlplus 格式化命令
        upper_strip = sql.upper().lstrip()
        for prefix in self._SKIP_PREFIXES:
            if upper_strip.startswith(prefix):
                return None, None, None

        # 去掉末尾分号（Oracle 不接受语句末尾分号）
        sql_clean = sql.rstrip(";/ \t\n")
        if not sql_clean:
            return None, None, None

        # 跳过纯 / 行（sqlplus 执行上一条语句的斜杠）
        if sql_clean.strip() == "/":
            return None, None, None

        try:
            cur = self.conn.cursor()
            # 设置 fetch array size 提升性能
            cur.arraysize = 200
            cur.execute(sql_clean)
            if cur.description:
                cols = [d[0] for d in cur.description]
                # Oracle LOB 字段转 str
                rows = []
                for row in cur.fetchmany(500):
                    new_row = []
                    for val in row:
                        if hasattr(val, "read"):          # LOB 对象
                            try:
                                val = val.read()
                            except Exception:
                                val = "<LOB>"
                        new_row.append(val)
                    rows.append(tuple(new_row))
                cur.close()
                return cols, rows, None
            cur.close()
            return None, None, None
        except UnicodeEncodeError as e:
            return None, None, f"[编码错误] SQL含非ASCII字符无法发送: {e}"
        except Exception as e:
            err_str = str(e)
            # 忽略某些常见的无权限/对象不存在错误，不视为致命
            innocuous = (
                "ORA-00942",  # table or view does not exist
                "ORA-01031",  # insufficient privileges
                "ORA-00904",  # invalid identifier
                "ORA-01219",  # database not open
                "ORA-06550",  # PL/SQL compilation error
            )
            for code in innocuous:
                if code in err_str:
                    return None, None, f"[跳过] {err_str[:120]}"
            return None, None, err_str[:200]


# ─── 连接器工厂 ──────────────────────────────────────────────────
DB_CONNECTORS = {
    "1": ("MySQL",        MySQLConnector),
    "2": ("PostgreSQL",   PostgreSQLConnector),
    "3": ("达梦(DM)",     DaMengConnector),
    "4": ("GoldenDB",     GoldenDBConnector),
    "5": ("Oracle",       OracleConnector),
}

# SQL 文件对应关系
DB_SQL_FILES = {
    "MySQL":       "mysql_master_inspection.sql",
    "PostgreSQL":  "postgresql_master_inspection.sql",
    "达梦(DM)":    "dameng_master_inspection.sql",
    "GoldenDB":    "goldendb_master_inspection.sql",
    "Oracle":      "oracle_master_inspection.sql",
}


# ─── Word 报告生成器 ─────────────────────────────────────────────
class WordReportGenerator:
    """生成单个Word巡检报告"""

    MAX_ROWS_PER_TABLE = 200   # 超过此行数截断
    MAX_COLS_WARN      = 12    # 超过此列数警告

    # 各数据库 × 各报告类型的专项运维建议
    # 结构: { db_type关键字 : { report_type : [建议列表] } }
    _ROUTINE_ADVICE: Dict[str, Dict[str, List[str]]] = {

        # ─────────────────────────────────────────────────────────
        # MySQL
        # ─────────────────────────────────────────────────────────
        "MySQL": {
            "健康报告": [
                "建议每日检查连接数使用率（max_connections），超过 70% 时评估连接池或扩容",
                "建议开启慢查询日志（slow_query_log=ON，long_query_time=1），定期分析 mysqldumpslow",
                "建议监控活跃会话与长事务，超过 10 分钟未提交的事务需及时确认",
                "建议通过 mysqldump/xtrabackup 验证备份可恢复性，定期做恢复演练",
                "建议将本报告纳入日常自动化巡检流程（cron + 邮件告警）",
            ],
            "风险报告": [
                "建议对慢 SQL 执行 EXPLAIN，优先处理全表扫描（type=ALL）和 Using filesort",
                "建议对锁等待设置告警（innodb_lock_wait_timeout ≤ 30s），避免级联阻塞",
                "建议定期复核 SUPER/DBA 权限账户，遵循最小权限原则，删除闲置账户",
                "建议为无主键的大表补充主键，防止从库 row-based 复制性能劣化",
                "建议对死锁高发场景统一事务访问顺序，并开启 innodb_print_all_deadlocks",
            ],
            "参数异常": [
                "建议修改参数前先在测试库验证，通过 SET GLOBAL 动态生效后再写入 my.cnf",
                "建议 innodb_flush_log_at_trx_commit=1 且 sync_binlog=1 保障数据持久性",
                "建议 innodb_buffer_pool_size 设置为可用内存的 50%~70%",
                "建议使用 MySQLTuner 或 Percona Toolkit 工具定期复查参数合理性",
                "建议移除已废弃参数（deprecated），防止升级到新版本时启动失败",
            ],
            "空间风险": [
                "建议对磁盘/表空间使用率设置 75% 告警、85% 紧急告警阈值",
                "建议定期清理 binlog（expire_logs_days / binlog_expire_logs_seconds），防止磁盘打满",
                "建议对碎片率高的大表在业务低峰期执行 OPTIMIZE TABLE 或 ALTER TABLE ENGINE=InnoDB",
                "建议归档历史数据到冷存储，对分区表定期 DROP 已过期分区",
                "建议建立空间增长趋势监控，提前 30 天预测扩容需求",
            ],
            "HA风险": [
                "建议主从复制延迟超过 60 秒时触发告警，检查大事务/并行复制配置",
                "建议开启 GTID 复制（gtid_mode=ON），简化主从切换与故障恢复",
                "建议定期演练主从切换流程（MHA/Orchestrator），验证 RTO ≤ 30 秒",
                "建议 MGR 集群配置 VIP 漂移（keepalived/Proxy），节点故障时自动接管",
                "建议从库开启 read_only=ON / super_read_only=ON，防止误写",
            ],
        },

        # ─────────────────────────────────────────────────────────
        # PostgreSQL
        # ─────────────────────────────────────────────────────────
        "PostgreSQL": {
            "健康报告": [
                "建议每日检查连接数使用率（max_connections），超过 70% 时使用 PgBouncer 连接池",
                "建议开启 pg_stat_statements 扩展，定期分析 Top SQL 的执行计划",
                "建议监控长事务（age > 10min）与空闲事务（idle in transaction），及时处理",
                "建议通过 pg_basebackup 或 pgBackRest 定期验证备份可恢复性",
                "建议将本报告纳入日常自动化巡检流程（cron + 邮件告警）",
            ],
            "风险报告": [
                "建议对慢 SQL 执行 EXPLAIN (ANALYZE, BUFFERS)，优先处理 Seq Scan 大表",
                "建议设置 lock_timeout 和 statement_timeout，避免长时间锁等待阻塞业务",
                "建议定期复核 SUPERUSER/pg_read_all_data 等高权限，遵循最小权限原则",
                "建议对频繁更新/删除的大表定期执行 VACUUM ANALYZE，防止表膨胀",
                "建议检查并消除 Bloat 较高的表和索引，必要时 CLUSTER 或 REINDEX CONCURRENTLY",
            ],
            "参数异常": [
                "建议修改参数前先在测试库验证，reload 类参数用 pg_ctl reload 生效",
                "建议 shared_buffers 设置为物理内存的 25%，work_mem 按并发数合理分配",
                "建议 wal_level=replica 且 synchronous_commit=on 保障数据持久性",
                "建议定期使用 pgtune 工具重新评估参数，结合负载类型（OLTP/OLAP）调优",
                "建议移除已废弃参数，避免升级大版本（如 14→17）时出现启动报错",
            ],
            "空间风险": [
                "建议对磁盘/表空间使用率设置 75% 告警、85% 紧急告警阈值",
                "建议定期清理 WAL 归档目录（archive_cleanup_command），防止磁盘打满",
                "建议对膨胀率高的表执行 VACUUM FULL（业务低峰），索引执行 REINDEX CONCURRENTLY",
                "建议对历史归档分区表使用 DETACH PARTITION + DROP 释放磁盘空间",
                "建议建立表/索引大小增长监控，提前 30 天预测扩容需求",
            ],
            "HA风险": [
                "建议流复制延迟超过 60 秒时触发告警，检查网络带宽与 wal_keep_size 配置",
                "建议开启 pg_replication_slot，防止备库落后导致 WAL 被清理",
                "建议定期演练主从切换（Patroni/Repmgr），验证 RTO ≤ 30 秒",
                "建议从库设置 hot_standby=on，支持只读查询分流，同时开启 recovery_target_timeline='latest'",
                "建议对 Patroni 集群配置健康检查告警，节点 DCS（etcd/ZooKeeper）异常时自动告警",
            ],
        },

        # ─────────────────────────────────────────────────────────
        # Oracle
        # ─────────────────────────────────────────────────────────
        "Oracle": {
            "健康报告": [
                "建议每日检查告警日志（alert_<SID>.log），发现 ORA- 错误及时分析处理",
                "建议监控活跃会话与锁等待（V$SESSION / V$LOCK），超过 5 分钟需介入",
                "建议检查 RMAN 备份作业状态（V$RMAN_BACKUP_JOB_DETAILS），确保每日备份成功",
                "建议监控检查点频率与 Redo 日志切换速率，避免日志切换过快影响性能",
                "建议将本报告纳入日常自动化巡检流程，结合 OEM/Zabbix 实现告警联动",
            ],
            "风险报告": [
                "建议对 Top 等待 SQL 使用 DBMS_XPLAN.DISPLAY_CURSOR 分析执行计划，优先处理全表扫描",
                "建议对锁等待超过 30 秒的会话触发告警（监控 V$SESSION.SECONDS_IN_WAIT）",
                "建议定期复核 DBA 角色授权（DBA_ROLE_PRIVS），非必要账户及时回收",
                "建议对无索引外键（DBA_CONSTRAINTS type='R'）尽快补充索引，防止 TM 锁升级",
                "建议收集陈旧统计信息（DBMS_STATS.GATHER_STALE_STATS），防止执行计划走偏",
            ],
            "参数异常": [
                "建议修改参数前在测试库验证，动态参数用 ALTER SYSTEM SET … SCOPE=BOTH",
                "建议 SGA_TARGET + PGA_AGGREGATE_TARGET 合计不超过物理内存的 75%，防止 OOM",
                "建议 PROCESSES 和 SESSIONS 参数根据业务峰值预留 20% 余量，避免 ORA-00020",
                "建议定期使用 AWR/ADDM 报告评估参数合理性，优先处理 ADDM 的 FINDING",
                "建议检查并移除已废弃（ISDEPRECATED=TRUE）的隐含参数，防止升级失败",
            ],
            "空间风险": [
                "建议对所有表空间设置 75% 告警、85% 紧急告警（OEM 阈值或自定义脚本）",
                "建议定期清理 FRA 闪回区归档日志（RMAN DELETE ARCHIVELOG），防止 ORA-19809",
                "建议对行迁移率高（CHAIN_CNT/NUM_ROWS > 10%）的表执行 ALTER TABLE MOVE 重组",
                "建议对空置旧分区（NUM_ROWS=0 且 SIZE > 100MB）执行 DROP/TRUNCATE PARTITION",
                "建议通过 DBA_HIST_TBSPC_SPACE_USAGE 分析表空间增长趋势，提前 30 天预测扩容",
            ],
            "HA风险": [
                "建议监控 ADG 同步延迟（V$DATAGUARD_STATS apply lag），超过 60 秒触发告警",
                "建议每月至少执行一次 Data Guard Switchover 演练，验证备库数据完整性和 RTO",
                "建议检查归档传输间隙（V$ARCHIVE_GAP），发现 Gap 后立即手动注册缺失归档",
                "建议 RAC 集群监控节点互联延迟（GV$SYSTEM_EVENT gc cr request），AVG > 15ms 需排查",
                "建议 RAC 配置 ONS（Oracle Notification Service）实现节点故障时应用层自动 failover",
            ],
        },

        # ─────────────────────────────────────────────────────────
        # 达梦(DM)
        # ─────────────────────────────────────────────────────────
        "达梦(DM)": {
            "健康报告": [
                "建议每日检查 DM 系统日志（dm_YYYYMMDD.log），关注 ERROR/FATAL 级别记录",
                "建议监控活跃会话与长事务（V$SESSIONS / V$LONG_EXEC_SQLS），超过 10 分钟需介入",
                "建议通过 DM RMAN 或 dexp/dmrman 定期验证备份可恢复性，做恢复演练",
                "建议检查检查点和归档切换频率，避免归档日志堆积导致磁盘打满",
                "建议将本报告纳入日常自动化巡检流程，结合 DM 监控中心实现告警联动",
            ],
            "风险报告": [
                "建议对慢 SQL 使用 EXPLAIN 或 V$SQL_PLAN 分析执行计划，优先处理全表扫描",
                "建议对锁等待设置告警（监控 V$LOCK_WAIT），超过 30 秒的锁等待需立即介入",
                "建议定期复核 DBA 角色与高权限用户，删除无业务对应的闲置账户",
                "建议对无索引外键关系的大表补充索引，防止 DML 时发生锁升级",
                "建议统计信息过期（30 天未更新）的大表执行 UPDATE STATISTICS 刷新",
            ],
            "参数异常": [
                "建议修改 dm.ini 参数前先在测试库验证，动态参数可通过 SP_SET_PARA_VALUE 生效",
                "建议 BUFFER（数据缓冲区）设置为物理内存的 40%~60%，SORT_BUF_SIZE 按并发合理分配",
                "建议 ARCH_INI=1 开启归档，LOG_BUF_SIZE 不低于 4096（KB）保障 Redo 缓冲充足",
                "建议定期使用 DM AWR 或 DEM（达梦企业管理器）评估参数合理性",
                "建议检查已废弃或无效参数（PARA_TYPE=2），清理 dm.ini 中的冗余配置",
            ],
            "空间风险": [
                "建议对数据文件和表空间使用率设置 75% 告警、85% 紧急告警阈值",
                "建议定期清理归档日志（DMRMAN DELETE ARCHIVELOG），保留最近 7~14 天",
                "建议对碎片率高的大表执行 ALTER TABLE ... MOVE 或重建聚簇索引",
                "建议对历史数据分区表定期 DROP 已过期分区，释放磁盘空间",
                "建议通过 SYS_SPACE_USAGE 视图分析表空间增长趋势，提前规划扩容",
            ],
            "HA风险": [
                "建议实时主备（REALTIME 模式）同步延迟超过 30 秒时触发告警",
                "建议每季度执行一次主备切换演练（SWITCHOVER），验证备库可用性与 RTO",
                "建议检查 Redo 日志归档传输状态（V$ARCH_STATUS），发现断链立即处理",
                "建议 DM DSC 集群配置心跳网络冗余，避免单点心跳故障触发误切换",
                "建议从库开启只读模式（READ_ONLY=1），防止误操作写入备库",
            ],
        },

        # ─────────────────────────────────────────────────────────
        # GoldenDB
        # ─────────────────────────────────────────────────────────
        "GoldenDB": {
            "健康报告": [
                "建议每日检查各 CN/DN 节点告警日志，关注 ERROR 和 FATAL 级别事件",
                "建议监控 CN 协调节点与 DN 数据节点的连接数使用率，超过 70% 时评估扩容",
                "建议通过 GoldenDB 备份工具定期验证备份可恢复性，做全量恢复演练",
                "建议检查各分片数据分布均衡性，避免热点分片导致性能瓶颈",
                "建议将本报告纳入日常自动化巡检流程，结合 GoldenDB 管理平台实现告警联动",
            ],
            "风险报告": [
                "建议对跨分片慢查询使用 GoldenDB EXPLAIN 分析，优先优化全分片广播查询",
                "建议对分布式锁等待超过 30 秒的事务触发告警，避免级联阻塞",
                "建议定期复核高权限账户（SUPER/DBA），删除无业务对应的闲置账户",
                "建议避免使用跨分片 JOIN，改用应用层聚合或中间件路由优化",
                "建议对热点分片数据进行重分片（RESHARD），均衡各节点压力",
            ],
            "参数异常": [
                "建议修改参数前先在测试集群验证，CN/DN 节点参数需保持一致性",
                "建议 innodb_flush_log_at_trx_commit=1 且 sync_binlog=1 保障 DN 节点数据持久性",
                "建议各 DN 节点 innodb_buffer_pool_size 设置为各节点可用内存的 50%~70%",
                "建议定期使用 GoldenDB 诊断工具复查参数合理性，结合负载类型调优",
                "建议移除 CN/DN 节点上已废弃的 MySQL 参数，防止升级时启动失败",
            ],
            "空间风险": [
                "建议对各 DN 节点磁盘使用率设置 75% 告警、85% 紧急告警阈值",
                "建议定期清理各 DN 节点 binlog（expire_logs_days），防止单节点磁盘打满",
                "建议对碎片率高的大表在业务低峰期执行 OPTIMIZE TABLE 或重建分片表",
                "建议监控各分片数据量差异（skew ratio > 30% 需重新分片均衡）",
                "建议建立全节点空间增长趋势监控，提前 30 天预测集群整体扩容需求",
            ],
            "HA风险": [
                "建议主从复制延迟超过 60 秒时触发告警，检查大事务/并行复制配置",
                "建议开启 GTID 复制，简化 GoldenDB 主备切换与故障恢复流程",
                "建议定期演练节点故障切换流程，验证 CN/DN 故障时的自动 failover",
                "建议 GoldenDB 集群配置 VIP 漂移，CN 节点故障时自动接管连接入口",
                "建议从库 DN 节点开启 read_only=ON，防止应用误写非主节点",
            ],
        },
    }

    def __init__(self, report_type: str, db_type: str,
                 host: str, inspect_time: str, extra_info: str = ""):
        self.report_type   = report_type
        self.db_type       = db_type
        self.host          = host
        self.inspect_time  = inspect_time
        self.extra_info    = extra_info
        self.sections: List[Dict] = []   # 目录用（已执行的节）

        # ── 巡检统计，用于生成总结 ──────────────────────────────
        self._stats = {
            "critical_cnt":   0,                # CRITICAL 行总数
            "warning_cnt":    0,                # WARNING 行总数
            "total_rows":     0,                # 检查总数据行
            "issues": [],                       # [(severity, section_title, detail)]
        }

        self.doc = Document()
        self._setup_doc()

    # ── 文档基础设置 ────────────────────────────────────────────
    def _setup_doc(self):
        """页面、字体、段落样式"""
        section = self.doc.sections[0]
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        # A4 横版
        from docx.enum.section import WD_ORIENT
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = section.right_margin = Cm(1.8)
        section.top_margin  = section.bottom_margin = Cm(1.5)

        styles = self.doc.styles

        # 正文字体
        normal = styles["Normal"]
        normal.font.name = "微软雅黑"
        normal.font.size = Pt(9)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # Heading 1
        h1 = styles["Heading 1"]
        h1.font.name = "微软雅黑"
        h1.font.size = Pt(12)
        h1.font.bold = True
        h1._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # Heading 2
        h2 = styles["Heading 2"]
        h2.font.name = "微软雅黑"
        h2.font.size = Pt(10)
        h2.font.bold = True
        h2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── 封面 ────────────────────────────────────────────────────
    def _add_cover(self):
        doc = self.doc
        primary, bg = REPORT_THEME[self.report_type]
        icon = REPORT_ICON[self.report_type]

        # 顶部彩色横条（通过带底色的段落模拟）
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._shade_paragraph(header_para, primary)
        run = header_para.add_run(f"  {self.db_type} 数据库巡检报告  ")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size       = Pt(18)
        run.font.bold       = True
        run.font.name       = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        header_para.paragraph_format.space_before = Pt(4)
        header_para.paragraph_format.space_after  = Pt(4)

        doc.add_paragraph()

        # 大标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(f"{icon}  {self.report_type}")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        rgb = tuple(int(primary[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(*rgb)

        doc.add_paragraph()

        # 信息表格
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = "Table Grid"
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fields = [
            ("数据库类型", self.db_type),
            ("主机地址",   self.host),
            ("巡检时间",   self.inspect_time),
            ("报告类型",   self.report_type),
            ("备注",       self.extra_info or "—"),
        ]
        for i, (k, v) in enumerate(fields):
            row = info_table.rows[i]
            # 标签列
            cell_k = row.cells[0]
            self._set_cell_bg(cell_k, primary)
            p = cell_k.paragraphs[0]
            run = p.add_run(k)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(10)

            # 值列
            cell_v = row.cells[1]
            self._set_cell_bg(cell_v, bg)
            p2 = cell_v.paragraphs[0]
            run2 = p2.add_run(str(v))
            run2.font.name = "微软雅黑"
            run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run2.font.size = Pt(10)

        doc.add_page_break()

    # ── 着色工具 ────────────────────────────────────────────────
    @staticmethod
    def _shade_paragraph(para, fill_hex: str):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex.upper())
        pPr.append(shd)

    @staticmethod
    def _set_cell_bg(cell, fill_hex: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex.upper())
        tcPr.append(shd)

    # ── 章节标题 ────────────────────────────────────────────────
    def _add_section_heading(self, no: int, title: str, tag: str):
        label = f"[{tag}]" if tag else ""
        heading = self.doc.add_heading(f"§{no}. {label} {title}", level=1)
        heading.paragraph_format.space_before = Pt(8)

    # ── 数据表格 ────────────────────────────────────────────────
    def _add_result_table(self, columns: List[str], rows: List[tuple],
                          section_title: str = ""):
        """将查询结果写为Word表格，含风险色彩标注，同时统计风险行"""
        if not columns:
            return

        display_rows = rows[:self.MAX_ROWS_PER_TABLE]
        if len(rows) > self.MAX_ROWS_PER_TABLE:
            note = self.doc.add_paragraph(
                f"  ⚠ 结果共 {len(rows)} 行，仅展示前 {self.MAX_ROWS_PER_TABLE} 行"
            )
            note.runs[0].font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

        col_count = min(len(columns), self.MAX_COLS_WARN)
        columns   = columns[:col_count]

        tbl = self.doc.add_table(
            rows=1 + len(display_rows),
            cols=col_count
        )
        tbl.style = "Table Grid"

        # 表头
        hdr = tbl.rows[0]
        for i, col in enumerate(columns):
            cell = hdr.cells[i]
            self._set_cell_bg(cell, "1F3864")
            p = cell.paragraphs[0]
            run = p.add_run(str(col))
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 数据行
        sec_critical = 0
        sec_warning  = 0
        for r_idx, row_data in enumerate(display_rows):
            tr = tbl.rows[r_idx + 1]
            row_data = row_data[:col_count]

            # 判断整行风险等级
            row_str = " ".join(str(v) for v in row_data if v is not None)
            risk     = self._detect_risk(row_str)

            # 统计
            self._stats["total_rows"] += 1
            if risk == "CRITICAL":
                self._stats["critical_cnt"] += 1
                sec_critical += 1
            elif risk == "WARNING":
                self._stats["warning_cnt"] += 1
                sec_warning += 1

            for c_idx, val in enumerate(row_data):
                cell = tr.cells[c_idx]
                # 交替行底色
                bg = "F2F2F2" if (r_idx % 2 == 0) else "FFFFFF"
                if risk == "CRITICAL":
                    bg = "FDECEA"
                elif risk == "WARNING":
                    bg = "FFF9E6"
                self._set_cell_bg(cell, bg)

                p   = cell.paragraphs[0]
                run = p.add_run(self._fmt_value(val))
                run.font.size = Pt(8)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

                if risk == "CRITICAL":
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    run.font.bold = True
                elif risk == "WARNING":
                    run.font.color.rgb = RGBColor(0xD3, 0x54, 0x00)
                    run.font.bold = True

        # 将本节有问题的情况记录到 issues
        if section_title and sec_critical > 0:
            self._stats["issues"].append(
                ("CRITICAL", section_title,
                 f"发现 {sec_critical} 条严重(CRITICAL)记录，需立即处理")
            )
        if section_title and sec_warning > 0:
            self._stats["issues"].append(
                ("WARNING", section_title,
                 f"发现 {sec_warning} 条警告(WARNING)记录，建议尽快处理")
            )

    @staticmethod
    def _detect_risk(text: str) -> str:
        t = text.upper()
        if "[CRITICAL]" in t:
            return "CRITICAL"
        if "[WARNING]" in t or "[WARN]" in t:
            return "WARNING"
        return "OK"

    @staticmethod
    def _fmt_value(val) -> str:
        if val is None:
            return "NULL"
        if isinstance(val, (datetime.datetime, datetime.date)):
            return str(val)
        if isinstance(val, float):
            return f"{val:.4f}".rstrip("0").rstrip(".")
        return str(val)

    # ── 错误信息 ────────────────────────────────────────────────
    def _add_error(self, msg: str):
        p = self.doc.add_paragraph(f"  ⚠ 执行错误: {msg}")
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            p.runs[0].font.italic    = True
            p.runs[0].font.size      = Pt(8)

    # ── 空结果提示 ───────────────────────────────────────────────
    def _add_empty(self):
        p = self.doc.add_paragraph("  （无数据）")
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            p.runs[0].font.size      = Pt(8)

    # ── 添加节内容 ───────────────────────────────────────────────
    def add_section(self, section: Dict, results: List[Dict]):
        """写入一个巡检节及其查询结果"""
        self._add_section_heading(section["no"], section["title"], section["tag"])

        has_any = False
        for res in results:
            if res.get("error"):
                self._add_error(res["error"])
                has_any = True
            elif res.get("columns") and res.get("rows") is not None:
                self._add_result_table(
                    res["columns"], res["rows"],
                    section_title=section["title"]
                )
                has_any = True

        if not has_any:
            self._add_empty()

        self.doc.add_paragraph()   # 节间距

    # ── 巡检总结与建议 ──────────────────────────────────────────
    def _add_summary_to_doc(self):
        """在 self.doc 末尾追加"巡检总结与建议"章节"""
        doc      = self.doc
        stats    = self._stats
        primary, _ = REPORT_THEME[self.report_type]
        rgb_pri  = tuple(int(primary[i:i+2], 16) for i in (0, 2, 4))

        doc.add_page_break()

        # ── 章节标题 ────────────────────────────────────────────
        h = doc.add_heading("巡检总结与建议", level=1)
        h.paragraph_format.space_before = Pt(6)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(*rgb_pri)

        # ── 总体健康评分 ────────────────────────────────────────
        critical_cnt = stats["critical_cnt"]
        warning_cnt  = stats["warning_cnt"]
        total_rows   = stats["total_rows"]
        issues       = stats["issues"]

        # 计算得分：满分 100，每 CRITICAL -10，每 WARNING -3，最低 0
        score = max(0, 100 - critical_cnt * 10 - warning_cnt * 3)
        if score >= 90:
            grade, grade_color = "优良  ✅", "1F7A4D"
        elif score >= 70:
            grade, grade_color = "一般  💡", "D35400"
        else:
            grade, grade_color = "存在风险  ⚠️",  "C0392B"

        # 评分表格（2列）
        doc.add_heading("一、总体健康评估", level=2)
        score_tbl = doc.add_table(rows=4, cols=2)
        score_tbl.style = "Table Grid"
        score_items = [
            ("检查项总数",   f"{total_rows} 行数据"),
            ("严重问题数",   f"{critical_cnt} 项"),
            ("警告问题数",   f"{warning_cnt} 项"),
            ("综合健康得分", f"{score} / 100  —  {grade}"),
        ]
        for i, (k, v) in enumerate(score_items):
            row = score_tbl.rows[i]
            self._set_cell_bg(row.cells[0], "1F3864")
            p = row.cells[0].paragraphs[0]
            r = p.add_run(k)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True; r.font.size = Pt(9); r.font.name = "微软雅黑"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            bg = "FDECEA" if (i == 1 and critical_cnt > 0) else \
                 "FFF9E6" if (i == 2 and warning_cnt  > 0) else \
                 "E8F5E9" if score >= 90 else \
                 "FEF3E2" if score >= 70 else "FDECEA"
            self._set_cell_bg(row.cells[1], bg)
            p2 = row.cells[1].paragraphs[0]
            r2 = p2.add_run(v)
            r2.font.size = Pt(9); r2.font.name = "微软雅黑"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if i == 3:
                rgb_g = tuple(int(grade_color[j:j+2], 16) for j in (0, 2, 4))
                r2.font.color.rgb = RGBColor(*rgb_g)
                r2.font.bold = True

        doc.add_paragraph()

        # ── 问题清单 ───────────────────────────────────────────
        doc.add_heading("二、发现的问题清单", level=2)

        critical_issues = [(s, d) for sev, s, d in issues if sev == "CRITICAL"]
        warning_issues  = [(s, d) for sev, s, d in issues if sev == "WARNING"]

        if not issues:
            p = doc.add_paragraph("✅  本次巡检未发现异常，各指标均在正常范围内。")
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(0x1F, 0x7A, 0x4D)
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.name = "微软雅黑"
                p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        else:
            if critical_issues:
                ci_h = doc.add_paragraph("🔴  严重问题（需立即处理）")
                if ci_h.runs:
                    ci_h.runs[0].font.bold = True
                    ci_h.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                    ci_h.runs[0].font.size = Pt(9)
                    ci_h.runs[0].font.name = "微软雅黑"
                    ci_h.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                for idx, (sec_title, detail) in enumerate(critical_issues, 1):
                    p = doc.add_paragraph(f"    {idx}. 【{sec_title}】 {detail}")
                    if p.runs:
                        p.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.name = "微软雅黑"
                        p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                doc.add_paragraph()

            if warning_issues:
                wi_h = doc.add_paragraph("🟡  警告问题（建议尽快处理）")
                if wi_h.runs:
                    wi_h.runs[0].font.bold = True
                    wi_h.runs[0].font.color.rgb = RGBColor(0xD3, 0x54, 0x00)
                    wi_h.runs[0].font.size = Pt(9)
                    wi_h.runs[0].font.name = "微软雅黑"
                    wi_h.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                for idx, (sec_title, detail) in enumerate(warning_issues, 1):
                    p = doc.add_paragraph(f"    {idx}. 【{sec_title}】 {detail}")
                    if p.runs:
                        p.runs[0].font.color.rgb = RGBColor(0xD3, 0x54, 0x00)
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.name = "微软雅黑"
                        p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        doc.add_paragraph()

        # ── 运维建议 ───────────────────────────────────────────
        doc.add_heading("三、运维建议", level=2)
        # 先按 db_type 取，再按 report_type 取；两层都找不到给空列表
        db_key = self.db_type  # e.g. "Oracle", "MySQL", "达梦(DM)", "GoldenDB", "PostgreSQL"
        advices = (
            self._ROUTINE_ADVICE
            .get(db_key, {})
            .get(self.report_type, [])
        )
        if not advices:
            p = doc.add_paragraph("  暂无针对本数据库类型的专项建议。")
            if p.runs:
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.name = "微软雅黑"
                p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        else:
            for idx, advice in enumerate(advices, 1):
                p = doc.add_paragraph(f"  {idx}. {advice}")
                if p.runs:
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.name = "微软雅黑"
                    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        doc.add_paragraph()

        # ── 报告结尾 ───────────────────────────────────────────
        end_p = doc.add_paragraph(
            f"—— 报告生成时间：{self.inspect_time}  ·  "
            f"{self.db_type} {self.report_type}  ·  本报告由自动化巡检脚本生成，仅供参考 ——"
        )
        end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if end_p.runs:
            end_p.runs[0].font.size  = Pt(8)
            end_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            end_p.runs[0].font.italic    = True
            end_p.runs[0].font.name      = "微软雅黑"
            end_p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── 保存（修复：不再向 self.doc 写封面/目录，避免重复）──────
    def save(self, filepath: str) -> str:
        # 先追加总结章节到正文内容尾部
        self._add_summary_to_doc()

        # 新建最终文档：封面 → 目录 → 正文（含总结）
        final_doc = Document()
        self._rebuild_to(final_doc)
        self._add_footer_to(final_doc)
        final_doc.save(filepath)
        return filepath

    def _rebuild_to(self, new_doc: Document):
        """重建文档：封面 → 目录 → 内容"""
        # 页面设置
        sec = new_doc.sections[0]
        from docx.enum.section import WD_ORIENT
        sec.orientation  = WD_ORIENT.LANDSCAPE
        sec.page_width   = Cm(29.7)
        sec.page_height  = Cm(21.0)
        sec.left_margin  = sec.right_margin = Cm(1.8)
        sec.top_margin   = sec.bottom_margin = Cm(1.5)

        styles = new_doc.styles
        normal = styles["Normal"]
        normal.font.name = "微软雅黑"
        normal.font.size = Pt(9)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        primary, bg = REPORT_THEME[self.report_type]
        icon        = REPORT_ICON[self.report_type]

        # ─ 封面 ─
        hp = new_doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._shade_paragraph(hp, primary)
        run = hp.add_run(f"  {self.db_type} 数据库综合巡检报告  ")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(16); run.font.bold = True; run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        new_doc.add_paragraph()
        tp = new_doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = tp.add_run(f"{icon}  {self.report_type}")
        run2.font.size = Pt(26); run2.font.bold = True; run2.font.name = "微软雅黑"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        rgb = tuple(int(primary[i:i+2], 16) for i in (0, 2, 4))
        run2.font.color.rgb = RGBColor(*rgb)
        new_doc.add_paragraph()

        # 信息框
        it = new_doc.add_table(rows=5, cols=2)
        it.style = "Table Grid"
        it.alignment = WD_TABLE_ALIGNMENT.CENTER
        fields = [
            ("数据库类型", self.db_type),
            ("主机地址",   self.host),
            ("巡检时间",   self.inspect_time),
            ("报告类型",   self.report_type),
            ("备注",       self.extra_info or "—"),
        ]
        for i, (k, v) in enumerate(fields):
            rw = it.rows[i]
            self._set_cell_bg(rw.cells[0], primary)
            p = rw.cells[0].paragraphs[0]
            r = p.add_run(k)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.bold = True
            r.font.name = "微软雅黑"; r.font.size = Pt(10)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            self._set_cell_bg(rw.cells[1], bg)
            p2 = rw.cells[1].paragraphs[0]
            r2 = p2.add_run(str(v))
            r2.font.name = "微软雅黑"; r2.font.size = Pt(10)
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        new_doc.add_page_break()

        # ─ 目录 ─
        toc_h = new_doc.add_paragraph("本报告目录")
        toc_h.runs[0].font.bold = True; toc_h.runs[0].font.size = Pt(12)
        toc_h.runs[0].font.name = "微软雅黑"
        toc_h.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        for s in self.sections:
            ip = new_doc.add_paragraph(f"    §{s['no']}  [{s['tag']}]  {s['title']}" if s['tag']
                                       else f"    §{s['no']}  {s['title']}")
            ip.paragraph_format.space_after = Pt(2)
            if ip.runs:
                ip.runs[0].font.size = Pt(9)
                ip.runs[0].font.name = "微软雅黑"
                ip.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        new_doc.add_page_break()

        # ─ 正文 ─ 将 self.doc 的段落/表格复制到 new_doc
        from copy import deepcopy
        for element in self.doc.element.body:
            new_doc.element.body.append(deepcopy(element))

    def _add_footer_to(self, doc: Document):
        sec    = doc.sections[0]
        footer = sec.footer
        para   = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(
            f"{self.db_type} · {self.report_type} · {self.inspect_time}  |  本报告由自动化巡检生成，仅供参考"
        )
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


# ─── 主巡检类 ────────────────────────────────────────────────────
class DBInspector:
    def __init__(self):
        self.connector:   Optional[BaseConnector] = None
        self.sql_text:    str = ""
        self.sections:    List[Dict] = []
        self.inspect_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.timestamp    = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        self.progress     = ProgressTracker()
        # report_type → WordReportGenerator
        self.generators: Dict[str, WordReportGenerator] = {}

    # ── 步骤1: 选择数据库类型 ────────────────────────────────────
    def choose_db_type(self):
        print("\n" + "=" * 55)
        print("  多数据库综合巡检工具  v1.1")
        print("=" * 55)
        print("  请选择数据库类型:")
        for k, (name, _) in DB_CONNECTORS.items():
            print(f"    {k}. {name}")
        choice = input("\n  输入编号 [1]: ").strip() or "1"
        if choice not in DB_CONNECTORS:
            print("  无效选择，默认 MySQL")
            choice = "1"
        name, cls = DB_CONNECTORS[choice]
        self.connector = cls()
        print(f"\n  已选择: {name}")

    # ── 步骤2: 加载 SQL 文件 ─────────────────────────────────────
    def load_sql_file(self, sql_file: Optional[str] = None) -> bool:
        db_name = self.connector.db_type
        if not sql_file:
            sql_file = DB_SQL_FILES.get(db_name, "")

        # 搜索路径
        search_dirs = [
            ".",
            os.path.dirname(os.path.abspath(__file__)),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "sql"),
        ]
        found = None
        for d in search_dirs:
            fp = os.path.join(d, sql_file)
            if os.path.exists(fp):
                found = fp
                break

        if not found:
            # 让用户手动输入路径
            print(f"\n  ⚠ 未找到SQL文件: {sql_file}")
            manual = input(f"  请输入SQL文件路径 (留空跳过): ").strip()
            if manual and os.path.exists(manual):
                found = manual
            else:
                print("  跳过SQL加载，将生成空报告")
                return False

        with open(found, "r", encoding="utf-8", errors="ignore") as f:
            self.sql_text = f.read()

        print(f"  ✓ 已加载SQL文件: {found}  ({len(self.sql_text)} bytes)")
        return True

    # ── 步骤3: 解析SQL为节 ───────────────────────────────────────
    def parse_sections(self):
        parser = SqlSectionParser(self.sql_text, self.connector.db_type)
        self.sections = parser.parse()
        print(f"  ✓ 解析到 {len(self.sections)} 个巡检节")

        # 按报告分类统计
        cats: Dict[str, int] = {}
        for s in self.sections:
            cats[s["category"]] = cats.get(s["category"], 0) + 1
        for cat, cnt in cats.items():
            print(f"      {cat}: {cnt} 节")

    # ── 步骤4: 初始化报告生成器 ──────────────────────────────────
    def init_generators(self):
        db_name = self.connector.db_type
        host    = f"{self.connector.host}:{self.connector.port}"
        for rtype in REPORT_TYPES:
            gen = WordReportGenerator(
                report_type=rtype,
                db_type=db_name,
                host=host,
                inspect_time=self.inspect_time,
            )
            self.generators[rtype] = gen

    # ── 步骤5: 执行巡检并填充报告 ───────────────────────────────
    def run_inspection(self):
        total = len(self.sections)
        pbar  = self.progress.start_step("执行SQL巡检", total)

        for sec in self.sections:
            cat = sec["category"]
            gen = self.generators.get(cat)
            if not gen:
                continue

            results = []
            for sql in sec.get("sqls", []):
                if not sql.strip():
                    continue
                cols, rows, err = self.connector.execute_query(sql)
                if err:
                    results.append({"error": err})
                elif cols:
                    results.append({"columns": cols, "rows": rows or []})

            gen.add_section(sec, results)
            gen.sections.append(sec)   # 供目录使用

            self.progress.tick(pbar, f"§{sec['no']} {sec['title'][:30]}")

        self.progress.finish_step(pbar)

    # ── 步骤6: 生成 Word 文件 ────────────────────────────────────
    def generate_reports(self, output_dir: str = ".") -> List[str]:
        os.makedirs(output_dir, exist_ok=True)
        files = []
        db_safe = re.sub(r"[^\w\u4e00-\u9fff]", "_", self.connector.db_type)
        host_safe = re.sub(r"[^\w]", "_", self.connector.host)

        pbar = self.progress.start_step("生成Word报告", len(REPORT_TYPES))

        for rtype in REPORT_TYPES:
            gen   = self.generators[rtype]
            fname = f"{db_safe}_{host_safe}_{rtype}_{self.timestamp}.docx"
            fpath = os.path.join(output_dir, fname)
            gen.save(fpath)
            files.append(fpath)
            self.progress.tick(pbar, f"已生成: {fname}")
            print(f"  📄 {rtype}: {fpath}")

        self.progress.finish_step(pbar)
        return files

    # ── 主流程 ───────────────────────────────────────────────────
    def run(self):
        try:
            # 1. 选择数据库类型
            self.choose_db_type()

            # 2. 获取连接信息
            self.connector.get_connection_info()

            # 3. 连接数据库
            pbar = self.progress.start_step("连接数据库")
            connected = self.connector.connect()
            self.progress.finish_step(pbar)

            # 4. 加载SQL文件
            pbar = self.progress.start_step("加载SQL文件")
            self.load_sql_file()
            self.progress.finish_step(pbar)

            # 5. 解析SQL节
            pbar = self.progress.start_step("解析SQL节")
            self.parse_sections()
            self.progress.finish_step(pbar)

            # 6. 初始化报告生成器
            self.init_generators()

            # 7. 执行巡检
            if connected and self.sections:
                self.run_inspection()
            elif not connected:
                print("\n  ⚠ 数据库未连接，将生成空白报告（仅含封面与目录）")
                for sec in self.sections:
                    cat = sec["category"]
                    gen = self.generators.get(cat)
                    if gen:
                        gen.add_section(sec, [])
                        gen.sections.append(sec)

            # 8. 生成报告
            output_dir = input("\n  报告输出目录 [./reports]: ").strip() or "./reports"
            files = self.generate_reports(output_dir)

            # 汇总
            print(f"\n{'=' * 55}")
            print(f"  🎉 巡检完成！共生成 {len(files)} 份Word报告")
            for f in files:
                print(f"    · {f}")
            print(f"\n  ⏱ 总耗时: {time.time() - self.progress.start_time:.2f}s")
            self.progress.summary()

        except KeyboardInterrupt:
            print("\n\n⚡ 用户中断")
        except Exception as e:
            print(f"\n💥 程序错误: {e}")
            traceback.print_exc()
        finally:
            if self.connector:
                self.connector.close()
                print("  ✓ 数据库连接已关闭")


# ─── 命令行入口 ──────────────────────────────────────────────────
def main():
    print("=" * 55)
    print("  多数据库综合巡检报告生成工具  v1.1")
    print("  支持: MySQL / PostgreSQL / 达梦(DM) / GoldenDB / Oracle")
    print("  输出: 5份Word报告 (健康/风险/参数/空间/HA)")
    print("=" * 55)

    # 快捷模式: python db_inspection.py --demo [db_type]
    # db_type: mysql(默认) / pg / dm / goldendb / oracle
    demo_mode = "--demo" in sys.argv
    if demo_mode:
        print("\n  [演示模式] 不连接数据库，直接解析SQL生成空格式报告")

        # 根据命令行参数选择演示数据库类型
        demo_db = "mysql"
        for arg in sys.argv[1:]:
            if arg.lower() in ("oracle", "ora"):
                demo_db = "oracle"
            elif arg.lower() in ("pg", "postgresql", "postgres"):
                demo_db = "postgresql"
            elif arg.lower() in ("dm", "dameng"):
                demo_db = "dm"
            elif arg.lower() in ("goldendb", "golden"):
                demo_db = "goldendb"

        demo_map = {
            "mysql":      (MySQLConnector,      "mysql_master_inspection.sql",      "demo-mysql-host",      3306),
            "postgresql": (PostgreSQLConnector,  "postgresql_master_inspection.sql", "demo-pg-host",         5432),
            "dm":         (DaMengConnector,       "dameng_master_inspection.sql",     "demo-dm-host",         5236),
            "goldendb":   (GoldenDBConnector,     "goldendb_master_inspection.sql",   "demo-goldendb-host",   3308),
            "oracle":     (OracleConnector,       "oracle_master_inspection.sql",     "demo-oracle-host",     1521),
        }
        cls, sql_file, demo_host, demo_port = demo_map[demo_db]

        inspector = DBInspector()
        inspector.connector = cls()
        inspector.connector.host = demo_host
        inspector.connector.port = demo_port
        inspector.connector.user = "demo"

        pbar = inspector.progress.start_step("加载SQL文件")
        inspector.load_sql_file(sql_file)
        inspector.progress.finish_step(pbar)

        pbar = inspector.progress.start_step("解析SQL节")
        inspector.parse_sections()
        inspector.progress.finish_step(pbar)

        inspector.init_generators()
        for sec in inspector.sections:
            cat = sec["category"]
            gen = inspector.generators.get(cat)
            if gen:
                gen.add_section(sec, [{"error": "演示模式，未执行SQL"}])
                gen.sections.append(sec)

        output_dir = f"./reports/{demo_db}"
        inspector.generate_reports(output_dir)
        inspector.progress.summary()
    else:
        inspector = DBInspector()
        inspector.run()


if __name__ == "__main__":
    main()
